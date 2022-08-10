'''


'''

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from pathlib import Path
from nrotc.models import Members
   
def getUserVariables():
    obj = Members.objects.last()
    variables = {
        'origin_code':getattr(obj, 'origin_code'), 
        'serial_num':getattr(obj, 'serial_num'), 
        'letter_date':getattr(obj, 'letter_date'), 
        'unit':getattr(obj, 'unit'), 
        'student_rank':getattr(obj, 'student_rank'), 
        'student_name':getattr(obj, 'student_name'), 
        'student_service':getattr(obj, 'student_service'), 
        'trigger':getattr(obj, 'trigger'), 
        'senior_member':getattr(obj, 'senior_rank') + ' ' + getattr(obj, 'senior_name') + ' ' + getattr(obj, 'senior_service'), 
        'member_1':getattr(obj, 'member1_rank') + ' ' + getattr(obj, 'member1_name') + ' ' + getattr(obj, 'member1_service'),
        'member_2':getattr(obj, 'member2_rank') + ' ' + getattr(obj, 'member2_name') + ' ' + getattr(obj, 'member2_service'),
        'recorder':getattr(obj, 'recorder_rank') + ' ' + getattr(obj, 'recorder_name') + ' ' + getattr(obj, 'recorder_service'), 

        'time':getattr(obj, 'prb_time'), 
        'date':getattr(obj, 'prb_date'), 
        'location':getattr(obj, 'prb_location'), 
        'student_uniform':getattr(obj, 'student_uniform'), 
        'CO_name':getattr(obj, 'CO_name'), 
        'interim_LOA':getattr(obj, 'student_on_LOA'),

        #prb report additions
        'present':getattr(obj, 'present'),
        'university':getattr(obj, 'university'),
        'major':getattr(obj, 'major'),
        'gpa':getattr(obj, 'gpa'),
        'credits':getattr(obj, 'credits'),
        'aptitude':getattr(obj, 'aptitude'),
        'academic_history':getattr(obj, 'academic_history'),
        'disciplinary_history':getattr(obj, 'disciplinary_history'),
        'performance_history':getattr(obj, 'performance_history'),
        'board_findings':getattr(obj, 'board_findings'),
        'board_rec':getattr(obj, 'board_rec'),
        }
    return variables

def createPRBConveningOrder():
    BASE_DIR = Path(__file__).resolve().parent.parent
    scripts_dir = str(BASE_DIR) + '/scripts/'
    variables = getUserVariables()

    document = Document(scripts_dir + 'template.docx')

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times'
    font.size = Pt(12)
    
    paragraph = document.add_paragraph('\t\t\t\t\t\t\t\t\t\t1533')
    paragraph.add_run('\n\t\t\t\t\t\t\t\t\t\tSer ' + variables['origin_code'])
    paragraph.add_run('/')
    paragraph.add_run(variables['serial_num'])
    paragraph.add_run('\n\t\t\t\t\t\t\t\t\t\t')
    paragraph.add_run(variables['letter_date'])
    paragraph.add_run('\n\nFrom:  Commanding Officer, Naval Reserve Officers Training Corps Unit, ')
    paragraph.add_run('\n\t')
    paragraph.add_run(variables['unit'])
    paragraph.add_run('\nTo:      ')
    paragraph.add_run(variables['student_rank'])
    paragraph.add_run(' ')
    paragraph.add_run(variables['student_name'])
    paragraph.add_run(', ')
    paragraph.add_run(variables['student_service'])
    paragraph.add_run('\n\nSubj:   PERFORMANCE REVIEW BOARD CONVENING ORDER')
    paragraph.add_run('\n\nRef:     (a) NSTC M-1533.2D')
    paragraph.add_run('''\n\n1.  Pursuant to reference (a), a performance review board (PRB) shall be convened to investigate and make recommendations on your ''')
    paragraph.add_run(variables['trigger'])
    paragraph.add_run('.')
    paragraph.add_run('\n\n2.  All PRB members shall be familiar with the applicable portions of reference (a). PRB members will be:\n\n\t')
    paragraph.add_run(variables['senior_member'])
    paragraph.add_run(', (Voting, Senior Member)')
    paragraph.add_run('\n\t')
    paragraph.add_run(variables['member_1'])
    paragraph.add_run(', (Voting)')
    paragraph.add_run('\n\t')
    paragraph.add_run(variables['member_2'])
    paragraph.add_run(', (Voting)')
    paragraph.add_run('\n\t')
    paragraph.add_run(variables['recorder'])
    paragraph.add_run(', (Non-voting, recorder)')
    paragraph.add_run('\n\n3.  The PRB will be convened at: ')
    paragraph.add_run(variables['time'] + ', ' + variables['date'] + ', ' + variables['location'] + '.')
    paragraph.add_run('\n\n4.  You may request to waive your right to a PRB if the reason for your PRB is Drop on Request, physical disqualification as documented by BUMED, or if the host institution has denied your continuation at the institution (not applicable for STA-21 or MECEP Students). The Professor of Naval Science may deny this request if a PRB is deemed necessary.\n\n5.  The PRB may recommend that any of the following actions be taken. Ramifications of these actions are described in reference (a).\n\n     a. No action.\n\n     b. Issuance of a 30-day compliance letter to investigate medical concerns.\n\n     c. Warning.\n\n     d. Probation.\n\n     e. Leave of Absence.\n\n     f. Disenrollment from the Naval Reserve Officers Training Corps (NROTC) Program. A recommendation of Directed Active Enlisted Service (DAES), recoupment/non-recoupment of scholarship funds may apply.')
    paragraph.add_run('\n\n6.  Contact your NROTC Class Advisor for a required counseling session prior to the PRB.')
    paragraph.add_run('\n\n7.  For midshipmen, your attendance at the PRB is strongly recommended, but not mandatory. Active duty students (STA-21 and MECEP) are required to attend. The uniform for the hearing is ')
    paragraph.add_run(variables['student_uniform'])
    paragraph.add_run('\n\n8.  You will be provided with a copy of the PRB Report no later than five business days after the PRB. If you choose to dispute or respond to the PRB’s findings, you must do so within five business days of receipt of the PRB Report.')
    paragraph.add_run('\n\n\n\n\t\t\t\t\t\t')
    paragraph.add_run(variables['CO_name'])
    run = paragraph.add_run()
    run.add_break(break_type=WD_BREAK.PAGE)
    paragraph = document.add_paragraph('From:  ')
    paragraph.add_run(variables['student_rank'])
    paragraph.add_run(' ')
    paragraph.add_run(variables['student_name'])
    paragraph.add_run(', ')
    paragraph.add_run(variables['student_service'])
    paragraph.add_run('\nTo:      Commanding Officer, Naval Reserve Officers Training Corps Unit, ')
    paragraph.add_run('\n\t')
    paragraph.add_run(variables['unit'])
    paragraph.add_run('\n\n1.  I acknowledge receipt of the PRB Convening Order and all documents that may be presented against me at the PRB.\n\n2.  I understand that I have the following rights with regard to these proceedings (student initials required for each optional statement):\n\n     a. The right to appear before the PRB at my own expense.\n\nI ____WILL/____WILL NOT appear before the PRB.\n\n     b. The right to challenge a member of this PRB for cause. The Senior Member of the PRB will make the final determination of a member’s suitability to serve on this PRB. If I wish to challenge the Senior Member of the PRB for cause, the PNS will determine the suitability.\n\nI ____DO/____DO NOT challenge a member of the PRB for cause.\n\n     c. The right to submit a written statement on my behalf.\n\nI ____WILL/____WILL NOT submit a written statement.\n\n     d. The right to present documents or witnesses on my behalf, at my own expense.\n\nI ____WILL/____WILL NOT present witnesses or documents to this PRB (Witness list and/or documents due to PNS at least 48 hours to the PRB).\n\n     e. The right to review, with my advisor, my personnel record.\n\nI ____DO/____DO NOT request to review, with my advisor, my personnel record.\n\n     f. The right to a five full business day waiting period to convene this PRB.\n\nI ____DO/____DO NOT waive my right to the five full business day waiting period to convene this PRB.\n\n     g. The right to retain counsel at my own expense.\n\nI ____WILL/____WILL NOT retain legal counsel for this PRB (Legal counsel information due to PNS at least 48 hours prior to the PRB).\n\n     h. The right to have observers attend the PRB.\n\nI ____WILL/____WILL NOT have observers at the PRB (must be identified at least 48 hours prior to the PRB).')
    
    if variables['interim_LOA'] == 'Yes':
        paragraph.add_run('\n\n3.  I acknowledge I am being placed on Interim Leave of Absence (LOA). I also understand that by being placed on LOA my subsistence allowance has been suspended and I am financially responsible for any required tuition and fees until my removal from LOA.')
    
    paragraph.add_run('\n\n\t\t\t\t\t_______________________________________')
    paragraph.add_run('\t\t\t\t\t\t  SIGNATURE                     DATE')

    run = paragraph.add_run()
    run.add_break(break_type=WD_BREAK.PAGE)
    table = document.add_table(rows=3, cols=1)
    table.style = 'TableGrid'

    cell = table.cell(0, 0)
    cell.text = '\n\t\t\t\t\t    NROTC\n\t\t\t     PERFORMANCE REVIEW BOARD\n\t\t\t          PRIVACY ACT STATEMENT'
    cell = table.cell(0, 1)
    cell.text = "Under the authority of the 5 U.S.C.A. sect. 301, 10 U.S.C.A. sect. 6011, U.S. Navy Regulations (articles 0802 and 0819) and NSTCNOTE 5210 information regarding your personal background may be requested in order to provide the performance review board in your case with additional information upon which to recommend your retention or separation from the NROTC program. The information provided by you will become a permanent part of the performance review board's record of proceedings and may be used by officials of the Department of the Navy in making recommendations or decisions in your case and by employees and officials of the Department of Defense, the Veterans' Administration and other Federal or State agencies in the performance of their official duties. You are not required to provide this information; however, failure to do so could result in the performance review board's discounting or giving little weight to any testimony which you may give and could deprive the board of valuable information on your behalf which it might otherwise consider in making a recommendation in your case."
    cell = table.cell(0, 2)
    cell.text = "\nI fully understand the privacy act statement listed above in regards to my Performance Review Board.\n\n\t\t\t\t_______________________________________\t\t\t\t\tSIGNATURE\t\t\t\tDATE"


    document.save(scripts_dir + 'PRB-Convening-Order.docx')

    return document

def createPRBReport():
    BASE_DIR = Path(__file__).resolve().parent.parent
    scripts_dir = str(BASE_DIR) + '/scripts/'
    variables = getUserVariables()

    document = Document(scripts_dir + 'template.docx')

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times'
    font.size = Pt(12)


    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = f'PERFORMANCE REVIEW BOARD REPORT ICO {variables["student_rank"]} {variables["student_name"]}, {variables["student_service"]}'
    
    paragraph = document.add_paragraph('\t\t\t\t\t\t\t\t\t\t1533')
    paragraph.add_run('\n\t\t\t\t\t\t\t\t\t\tSer ' + variables['origin_code'])
    paragraph.add_run('/')
    paragraph.add_run(variables['serial_log'])
    paragraph.add_run('\n\t\t\t\t\t\t\t\t\t\t')
    paragraph.add_run(variables['letter_date'])

    paragraph.add_run('\n\nMEMORANDUM')

    paragraph.add_run('\n\nFrom:  Senior Member, Naval Reserve Officers Training Corps Unit,\n\t' + variables['unit'])
    paragraph.add_run('\nTo:      Commanding Officer, Naval Reserve Officers Training Corps Unit,\n\t' + variables['unit'])

    paragraph.add_run('\n\nSubj:   PERFORMANCE REVIEW BOARD REPORT ICO ')
    paragraph.add_run(variables['student_rank'])
    paragraph.add_run(' ')
    paragraph.add_run(variables['student_name'])
    paragraph.add_run(', ')
    paragraph.add_run(variables['student_service'])

    paragraph.add_run('\n\nRef:     (a) NSTC M-1533.2D')

    paragraph.add_run('\n\nEncl:    (1) NSTC M-1533.2D')

    paragraph.add_run('\n\n1.  Per reference (a), the performance review board (PRB) directed by enclosure (1) convened at: ')
    paragraph.add_run(variables['prb_time'] + ', ' + variables['prb_date'] + ', ' + variables['prb_location'] + '.')
    
    paragraph.add_run('\n\n2.  Board members present:\n\n\t')
    paragraph.add_run(variables['senior_member'])
    paragraph.add_run(', (Voting, Senior Member)')
    paragraph.add_run('\n\t')
    paragraph.add_run(variables['member1'])
    paragraph.add_run(', (Voting)')
    paragraph.add_run('\n\t')
    paragraph.add_run(variables['member2'])
    paragraph.add_run(', (Voting)')
    paragraph.add_run('\n\t')
    paragraph.add_run(variables['recorder'])
    paragraph.add_run(', (Non-voting, recorder)')

    paragraph.add_run('\n\n3.  ' + variables['student_rank'] + ' ' + variables['student_name'] + ' ' + variables['present'] + ' present.')

    paragraph.add_run('\n\n4.  ' + variables['student_rank'] + ' Data:')
    paragraph.add_run('\n\n\tUniversity: ' + variables['university'])
    paragraph.add_run('\n\tMajor: ' + variables['major'])
    paragraph.add_run('\n\tCumulative GPA: ' + variables['gpa'])
    paragraph.add_run('\n\tCredits: ' + variables['credits'])
    paragraph.add_run('\n\tMilitary Aptitude: ' + variables['aptitude'])

    paragraph.add_run('\n\n5.  Academic History: ' + variables['academic_history'] + ' .')
    
    paragraph.add_run('\n\n6.  Disciplinary Problems: ' + variables['disciplinary_problems'] + ' .')

    paragraph.add_run('\n\n7.  Professional Performance History: ' + variables['performance_history'] + ' .')

    paragraph.add_run('\n\n8.  Board findings: By a vote of ' + variables['board_findings'] + ' .')

    paragraph.add_run('\n\n9.  Board recommendations: By a vote of ' + variables['board_rec'] + ' .')

    paragraph.add_run('\n\n\n\t\t\t\t\t\t')
    paragraph.add_run(variables['senior_sign'])

    paragraph.add_run('\n\n\n______________________      ______________________      ______________________')
    paragraph.add_run('\nVoting Member                         Voting Member                         Recorder')
    paragraph.add_run('\n\n------------------------------------------------------------------------------------------------------------')

    paragraph.add_run('From:  ')
    paragraph.add_run(variables['student_rank'])
    paragraph.add_run(' ')
    paragraph.add_run(variables['student_name'])
    paragraph.add_run(', ')
    paragraph.add_run(variables['student_service'])
    paragraph.add_run('\nTo:      Commanding Officer, Naval Reserve Officers Training Corps Unit, ')
    paragraph.add_run('\n\t')
    paragraph.add_run(variables['unit'])

    paragraph.add_run('\n\n1.  I have received the Board\'s recommendation.')

    paragraph.add_run('\n\n2.  I understand that I have five business days if I choose to submit a written statement in response to the Board\'s recommendation.')

    paragraph.add_run('\n\n3.  I _____WILL / _____WILL NOT submit a written statement in response to the Board\'s recommendation. (student initials required)')

    paragraph.add_run('\n\n\n\t\t\t\t\t_______________________________________')
    paragraph.add_run('\t\t\t\t\t\t  SIGNATURE                     DATE')

    document.save(scripts_dir + 'PRB-Report.docx')

    return document